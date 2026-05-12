from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\Law_frame_main")
OUT = ROOT / "artifacts" / "reports" / "lexframe_latest_changes_report_2026-05-10.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text: object, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(8.4)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc: Document, headers: list[str], rows: list[tuple], widths: list[float] | None = None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell_text(t.rows[0].cells[index], header, bold=True, color="FFFFFF")
        shade(t.rows[0].cells[index], "1F4E79")
    for row_index, row in enumerate(rows):
        cells = t.add_row().cells
        for index, value in enumerate(row):
            cell_text(cells[index], value)
            if index == 0:
                shade(cells[index], "EAF2F8" if row_index % 2 == 0 else "F7FBFE")
    if widths:
        for row in t.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph()
    return t


def note(doc: Document, title: str, body: str, fill: str = "EAF2F8") -> None:
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    shade(c, fill)
    p = c.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(31, 78, 121)
    p2 = c.add_paragraph(body)
    for run in p2.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(9)
    doc.add_paragraph()


def bullets(doc: Document, items: list[str | tuple[str, str]]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            a, b = item
            r = p.add_run(a + ": ")
            r.bold = True
            p.add_run(b)
        else:
            p.add_run(item)


def code(doc: Document, text: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    shade(c, "F2F2F2")
    p = c.paragraphs[0]
    for line in text.splitlines():
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8)
        p.add_run("\n")
    doc.add_paragraph()


tracked_files = [
    ("Backend / AI Gateway", "apps/backend/src/modules/ai-gateway/ai-gateway.service.ts", "+2 / -0", "Сервис AI Gateway: точечное расширение live chat path и safe metadata вокруг chat-provider поведения."),
    ("Backend / AI Gateway tests", "apps/backend/src/modules/ai-gateway/ai-provider.adapters.spec.ts", "+84 / -0", "Regression tests для CometAPI/OpenAI-compatible streaming request: endpoint, Authorization server-side, body keys, reasoning/thinking, SSE parsing."),
    ("Backend / AI Gateway", "apps/backend/src/modules/ai-gateway/ai-provider.adapters.ts", "+102 / -23", "Отдельный streaming chat completion path через backend-only fetch; SSE parser; provider-specific fields; no raw key in descriptors."),
    ("Backend / Chat tests", "apps/backend/src/modules/chat/chat-thread.service.spec.ts", "+56 / -1", "Проверки persistence user/assistant messages, failure event, отсутствие canned response, safe audit."),
    ("Backend / Chat", "apps/backend/src/modules/chat/chat-thread.service.ts", "+130 / -0", "Перевод project chat с structured/canned path на live AI Gateway streaming, persistence и failure handling."),
    ("Backend / Documents", "apps/backend/src/modules/documents/documents.controller.ts", "+48 / -0", "Новый endpoint upload bytes: POST /documents/:documentId/versions/:versionId/content."),
    ("Backend / Documents tests", "apps/backend/src/modules/documents/documents.service.spec.ts", "+94 / -0", "Tests для приёма реальных bytes, sha256, size mismatch, audit без raw content."),
    ("Backend / Documents", "apps/backend/src/modules/documents/documents.service.ts", "+87 / -1", "Валидация base64 content, size/hash/mime; audit document.version.content_received safe metadata."),
    ("Backend / Settings tests", "apps/backend/src/modules/settings/ai-settings.service.spec.ts", "+129 / -2", "Regression tests для AI connection test: /models не единственный PASS; chat completion visible-content readiness."),
    ("Backend / Settings", "apps/backend/src/modules/settings/ai-settings.service.ts", "+126 / -13", "Усиление AI settings/test connection flow, secret refs/fingerprints, route readiness и safe diagnostics."),
    ("Frontend / App shell tests", "apps/web/src/components/app-shell.test.tsx", "+21 / -0", "Regression: stale Activepieces browser token must be cleared outside embed route."),
    ("Frontend / App shell", "apps/web/src/components/app-shell.tsx", "+13 / -1", "Global cleanup Activepieces session tokens when leaving /automation embed route."),
    ("Frontend / Sidebar tests", "apps/web/src/components/shell/project-sidebar.test.tsx", "+34 / -5", "Tests for project chat list/sidebar updates and branch entries."),
    ("Frontend / Sidebar", "apps/web/src/components/shell/project-sidebar.tsx", "+50 / -40", "Project-scoped chat list instead of one-shot search; sidebar updates after chat stream."),
    ("Frontend / Upload tests", "apps/web/src/components/upload-dialog.test.tsx", "+36 / -6", "File selection + content upload test: File -> base64 -> backend content endpoint -> complete with sha256."),
    ("Frontend / Upload", "apps/web/src/components/upload-dialog.tsx", "+77 / -1", "Real file input, File metadata, byte upload before complete, query invalidation."),
    ("Frontend / Chat tests", "apps/web/src/features/ai-chat/components/LexFrameChatShell.test.tsx", "+15 / -0", "Regression: assistant marker scoped to assistant message, stream success/failure invalidates sidebar snapshots."),
    ("Frontend / Chat", "apps/web/src/features/ai-chat/components/LexFrameChatShell.tsx", "+29 / -1", "Stream completion/failure invalidates project chats/snapshot; UI running state controlled."),
    ("Frontend / Automation canvas", "apps/web/src/features/automation-canvas/activepieces-canvas-route.tsx", "+49 / -0", "Visible Запустить dry-run control on Activepieces iframe route; backend run endpoint used."),
    ("Frontend / Automation canvas tests", "apps/web/src/features/automation-canvas/activepieces-canvas-wrapper.test.tsx", "+32 / -0", "Regression: Activepieces browser JWT removed when iframe unmounts."),
    ("Frontend / Automation canvas", "apps/web/src/features/automation-canvas/activepieces-canvas-wrapper.tsx", "+2 / -0", "Wrapper cleanup calls shared Activepieces browser session cleanup."),
    ("Contracts/API", "packages/api-client/src/index.ts", "+13 / -0", "ApiClient method uploadDocumentVersionContent."),
    ("Contracts", "packages/contracts/src/domain.ts", "+13 / -0", "DocumentUploadContentRequest/Response contracts."),
    ("Contracts", "packages/contracts/src/errors/error-codes.ts", "+3 / -0", "DOCUMENT_UPLOAD_CONTENT_INVALID / HASH_MISMATCH / SIZE_MISMATCH."),
    ("Artifacts", "artifacts/stage17/playwright-browser-evidence.json", "+4 / -4", "Redacted old JWT-like browser evidence artifact."),
    ("Artifacts", "artifacts/stage19/chat-entrypoint-inventory.json", "+198 / -198", "Redacted old signed-url-like artifact pattern."),
    ("Artifacts", "artifacts/stage21/live-product-smoke-db-evidence.json", "+27 / -27", "Updated Stage21 DB evidence snapshot."),
    ("Artifacts", "artifacts/stage21/live-product-smoke-provider-probe.json", "+6 / -6", "Updated provider probe safe artifact."),
    ("Artifacts", "artifacts/stage21/live-product-smoke.json", "+33 / -33", "Updated live-product-smoke status artifact."),
    ("Frontend / Automation canvas", "apps/web/src/features/automation-canvas/activepieces-browser-session.ts", "NEW", "Shared safe cleanup helper for Activepieces-issued browser JWTs in local/session storage."),
    ("Frontend / Automation canvas tests", "apps/web/src/features/automation-canvas/activepieces-canvas-route.test.tsx", "NEW", "Regression: canvas route dry-run button calls LexFrame backend startAutomationRun with dry_run mode."),
    ("Artifacts", "artifacts/stage21/live-product-smoke-progress.jsonl", "NEW", "Stage21 progress evidence artifact."),
]

commands = [
    ("corepack pnpm --dir apps/web exec vitest run src/components/upload-dialog.test.tsx", "RED then PASS", "Первый запуск упал из-за отсутствия real file input. После фикса upload bytes test прошёл."),
    ("corepack pnpm --dir apps/web exec vitest run src/features/automation-canvas/activepieces-canvas-route.test.tsx", "RED then PASS", "Первый запуск упал из-за отсутствия кнопки Запустить dry-run. После фикса test прошёл."),
    ("corepack pnpm --dir apps/web exec vitest run src/features/automation-canvas/activepieces-canvas-wrapper.test.tsx", "RED then PASS", "Regression выявил оставшийся Activepieces JWT в storage; после cleanup фикса 6/6 tests PASS."),
    ("corepack pnpm --dir apps/web exec vitest run src/features/automation-canvas/activepieces-canvas-wrapper.test.tsx src/components/app-shell.test.tsx", "PASS", "9 tests PASS: wrapper cleanup + AppShell cleanup guard."),
    ("corepack pnpm --dir apps/backend exec jest src/modules/documents/documents.service.spec.ts --runInBand", "PASS", "10 tests PASS после добавления upload content validation."),
    ("corepack pnpm --dir apps/backend exec jest src/modules/activepieces/activepieces.service.spec.ts src/modules/activepieces/activepieces-session.service.spec.ts src/modules/documents/documents.service.spec.ts src/modules/chat/chat-thread.service.spec.ts --runInBand", "PASS", "4 suites / 22 tests PASS."),
    ("corepack pnpm --dir apps/web exec vitest run src/components/upload-dialog.test.tsx src/components/shell/project-automations-landing.test.tsx src/features/automation-canvas/use-activepieces-session.test.tsx src/features/automation-canvas/activepieces-canvas-route.test.tsx src/features/automation-canvas/activepieces-canvas-wrapper.test.tsx src/components/app-shell.test.tsx src/features/ai-chat/components/LexFrameChatShell.test.tsx", "PASS", "7 files / 22 tests PASS."),
    ("corepack pnpm --filter @lexframe/backend typecheck", "PASS", "Backend TypeScript noEmit PASS."),
    ("corepack pnpm --filter @lexframe/web typecheck", "PASS", "Web TypeScript noEmit PASS."),
    ("corepack pnpm --filter @lexframe/contracts typecheck", "PASS", "Contracts TypeScript noEmit PASS."),
    ("corepack pnpm --dir tests/e2e typecheck", "PASS", "E2E TypeScript noEmit PASS."),
    ("node scripts/stage21-up.mjs rebuild-backend", "PASS", "Backend image rebuilt/recreated; backend healthy; reverse proxy recreated."),
    ("node scripts/stage21-up.mjs rebuild-web", "PASS", "Web image rebuilt/recreated; Next build PASS; web healthy."),
    ("corepack pnpm secret-scan", "PASS", "Repo secret scan PASS."),
    ("Runtime artifact scan over artifacts + .playwright-mcp", "PASS after redaction", "providerKey=0, authBearer=0, bearerSk=0, JWT=0, signedUrl=0, privateKey=0."),
    ("docker exec ... DB evidence queries", "PASS", "workflow_runs=2 checked, documents=1 checked, chat_branches=1 checked, canvas_test_runs=1 checked, raw_secret_pattern_hits=0."),
]

bugs = [
    ("Project chat history not visible in sidebar", "Sidebar showed no chats or did not refresh after chat creation/stream.", "ProjectSidebar used one-shot searchChats({ q: '' }) and did not refetch project-scoped chats after stream lifecycle.", "Use project-scoped stage15 project chats and invalidate stage15-project-chats/stage15-project-snapshot after stream success/failure.", "Sidebar in Playwright shows existing project chats and branch entries without full reload."),
    ("Chat provider request path originally not using live streaming", "Project chat path was structured/non-streaming/canned-like and did not satisfy CometAPI streaming requirements.", "Normal chat was mixed with structured output route; request omitted streaming fields and provider-specific thinking metadata.", "Separate OpenAI-compatible streaming path, backend-only fetch, body contains stream/max_tokens/reasoning_effort/thinking, SSE parser handles reasoning chunks.", "Backend AI/chat targeted tests PASS; safe descriptors and provider probe artifacts updated."),
    ("CometAPI visible content empty after 200 OK", "Provider could return reasoning-only chunks and backend surfaced AI_GATEWAY_NOT_READY.", "Parser expected visible delta.content too narrowly and settings connection test treated /models as enough readiness.", "Broaden content extraction and add empty-content retry with visible-first params; settings readiness requires chat completion visible-content evidence.", "Settings/chat regression tests PASS."),
    ("Upload flow was metadata-only", "New upload completed an intent without real local file selection or bytes reaching backend.", "UploadDialog created intent then complete directly; no file input; no sha256 computed from actual content.", "Add file input, file-to-base64, backend content endpoint, content validation and sha256 audit metadata.", "Browser upload created document/version/storage rows; backend docs tests PASS."),
    ("Automation iframe route had no LexFrame run entry point", "Activepieces canvas opened, but LexFrame did not expose a run button over the iframe route.", "Run controls existed elsewhere, not on /automation embed route.", "Add Запустить dry-run control on canvas route that calls LexFrame backend startAutomationRun.", "Browser smoke created workflow_run and activepieces_run_binding rows."),
    ("Activepieces browser JWT persisted after leaving iframe", "sessionStorage contained an Activepieces JWT after navigation away from /automation.", "Embed runtime wrote token to browser storage; wrapper/AppShell did not clear stale AP token outside embed route.", "Shared cleanup helper detects Activepieces issuer JWT and removes it on wrapper unmount and outside embed route.", "Browser storage scan after navigation shows 0 hits; regression tests PASS."),
    ("Canvas branch live test cannot succeed on current workflow", "Branch test reaches backend but returns controlled RUNTIME_MAPPING_MISSING.", "Current workflow has manual_start -> end_success and no branch-capable runtime mapping/condition router node.", "No forced mutation of user workflow; controlled failure kept with audit evidence. True/false branch execution needs a branch-capable workflow fixture.", "canvas_test_runs row and audit events exist; status remains partial for live canvas branch behavior only."),
]

evidence = [
    ("Automation dry-run", "app.workflow_runs", "45ec0f45-d43d-4c9f-aada-3ab095161692", "mode=dry_run, status=queued, external ref stored as ap_run_45ec0f45d43d."),
    ("Automation binding", "app.activepieces_run_bindings", "workflow_run_id=45ec0f45...", "callback_token_hash and scoped_runtime_token_jti_hash present; raw tokens not persisted."),
    ("Runtime token", "app.activepieces_runtime_tokens", "workflow_run_id=45ec0f45...", "token_hash=true, jti_hash=true, expires_at recorded."),
    ("Automation audit", "audit.audit_events", "activepieces.run.started / activepieces.run.dispatched", "Safe metadata: automationId, mode, external refs, runtimeFlowId/runtimeProjectId; no raw JWT/header."),
    ("Document upload", "app.documents", "da52868e-d0b1-46ba-8f9f-4426b95924f5", "status=ready, current_version_id=b4e9370d..."),
    ("Document version", "app.document_versions", "b4e9370d-949a-4b9c-9fc2-5ae1196d4967", "original_filename=lexframe-upload-smoke.txt, mime=text/plain, size=21, sha256 stored."),
    ("Document storage", "app.document_storage_objects", "2 rows", "original + extracted_text storage object metadata created."),
    ("Document audit", "audit.audit_events", "document.upload.intent_created / content_received / completed", "Only safe metadata: documentId, versionId, mimeType, sizeBytes, sha256."),
    ("Chat branch", "app.chat_thread_branches", "branch_thread_id=e1f31186-a888-40f3-8dfe-517602fe7ee2", "Lineage persisted for project_claim_001; sidebar route visible."),
    ("Canvas branch test", "app.canvas_test_runs", "66643d92-9600-4d0d-903a-8f1316ace112", "Controlled failure: RUNTIME_MAPPING_MISSING, audit events written."),
    ("Secret DB scan", "DB text/json evidence", "raw_secret_pattern_hits=0", "No sk-pattern, Authorization Bearer, Bearer sk-, or JWT-shaped values in checked DB evidence areas."),
]

report_history = [
    ("Initial baseline/discovery", "I reported that settings persistence, AI settings write-only behavior, DB evidence and secret scan were confirmed on controlled failure path, while chat happy path remained blocked by provider/config.", "This framed independent domains: settings, chat/provider, projects, automations/runtime, secret-safety."),
    ("Domain-separated Stage21 gate", "I reported domain statuses such as Settings PASS, Projects PASS, Automations + runtime/canvas PASS, Secret-safety PASS, Chat + AI provider PARTIAL_EXTERNAL_PROVIDER_BLOCKER.", "Rule: do not block automation/settings/projects on chat provider failure; do not claim full PASS until all critical domains pass."),
    ("Credential source recovery", "I reported that request assembly was no longer the main problem once the safe descriptor proved endpoint/body/header shape; a stale local credential source remained a blocker until expected fingerprint/length appeared.", "Raw key was never printed. Fingerprints and sanitized lengths were used as evidence."),
    ("Project chat runtime plan", "I accepted the screenshot diagnosis: sidebar did not refresh chats and provider could return empty visible content. I planned retry-on-empty-visible-content, stronger settings readiness, and sidebar invalidation.", "Chat fix stayed backend-only for provider calls and kept frontend out of security decisions."),
    ("Runtime automation/upload QA final report", "I reported that automation dry-run and upload worked and persisted, but canvas branch live behavior stayed partial because the current workflow had no branch-capable node/runtime mapping.", "I also reported the security fix: Activepieces browser JWT was found in sessionStorage and then removed by wrapper/AppShell cleanup."),
]


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)
styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(9.5)
for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Aptos Display"
styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
styles["Heading 2"].font.color.rgb = RGBColor(46, 92, 122)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LexFrame")
r.bold = True
r.font.name = "Aptos Display"
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(31, 78, 121)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Подробный отчёт о последних изменениях в проекте")
r.font.name = "Aptos Display"
r.font.size = Pt(20)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Сформировано по истории текущего чата, diff рабочей копии, runtime evidence, DB/audit evidence и результатам проверок")
r.italic = True
r.font.size = Pt(10)

table(doc, ["Поле", "Значение"], [
    ("Дата отчёта", "10 мая 2026"),
    ("Рабочая директория", str(ROOT)),
    ("Ветка", "main"),
    ("Последний commit", "15003ed Доведён Stage21 live smoke до PASS"),
    ("Состояние рабочей копии", "Есть незакоммиченные изменения, созданные в рамках текущего набора работ и ранее в этом чате."),
    ("Формат", "DOCX / Word document"),
], [2.2, 5.6])
note(doc, "Важно о секретах", "В документ намеренно не включены raw provider keys, JWT, Authorization headers, signed URLs или приватные ключи. Все credential evidence сведены к fingerprints, lengths, статусам scan и безопасным metadata.", "FFF2CC")

doc.add_page_break()
doc.add_heading("Содержание", level=1)
bullets(doc, [
    "1. Executive summary",
    "2. Контекст и инварианты",
    "3. Хронология работы и история отчётов в чате",
    "4. Изменения по доменам",
    "5. Подробный файл-за-файлом обзор",
    "6. Баги, root cause и fixes",
    "7. Runtime evidence: browser, DB, audit",
    "8. Команды и результаты",
    "9. Secret-safety и redaction",
    "10. Оставшиеся ограничения и следующий шаг",
    "11. Приложения",
])

doc.add_heading("1. Executive summary", level=1)
for paragraph in [
    "В рамках текущего чата работа шла как инженерный release-gate, а не как визуальная проверка. Основные направления: довести Chat + AI provider path до реального backend-only CometAPI/OpenAI-compatible streaming поведения, устранить проблемы истории чатов и sidebar, проверить automation runtime/canvas, добавить run entry point, проверить branches и загрузку файлов, а также подтвердить persistence/audit/secret-safety evidence.",
    "На последнем проходе через Playwright было подтверждено: Activepieces iframe открывается; dry-run automation запускается через LexFrame backend и создаёт строки в product DB; upload теперь выбирает реальный локальный файл, отправляет bytes в backend, считает sha256, завершает upload contract и создаёт document/version/storage/audit evidence; browser direct external provider/runtime requests не обнаружены; storage cleanup после выхода из iframe удаляет Activepieces JWT.",
    "Самый существенный новый security bug, найденный в этом проходе: Activepieces embed runtime оставлял browser JWT в sessionStorage после ухода с iframe. Это было исправлено regression-test-first: сначала тест показал failure, затем добавлен shared cleanup helper и AppShell guard, после чего тесты и browser storage scan стали зелёными.",
    "По canvas branch live execution есть контролируемое ограничение: branch test через UI доходит до backend и пишет canvas_test_runs/audit, но текущий workflow содержит manual_start -> end_success и не имеет branch-capable runtime mapping. Система вернула controlled RUNTIME_MAPPING_MISSING, не raw crash и не секретную ошибку. Я не стал насильно мутировать пользовательский workflow для искусственного PASS.",
]:
    doc.add_paragraph(paragraph)
note(doc, "Итог последнего прохода", "Automation run, file upload, chat branch sidebar visibility, DB/audit evidence, typecheck, targeted tests и secret scan прошли. Единственное ограничение: live true/false canvas branch execution требует branch-capable workflow fixture.", "E2F0D9")

doc.add_heading("2. Контекст и проектные инварианты", level=1)
bullets(doc, [
    ("Source of truth", "backend и product DB остаются владельцами пользователей, workspace, проектов, чатов, сообщений, настроек, автоматизаций, runtime evidence и audit."),
    ("Frontend security", "frontend не принимает security decisions, не хранит provider secrets и не вызывает AI provider напрямую."),
    ("Provider/runtime", "внешние AI/runtime сервисы используются только через backend gateway/proxy/runtime контур."),
    ("Secrets", "raw provider key, JWT, service secrets, Authorization headers и signed URLs не должны попадать в код, fixtures, docs, artifacts, DB plaintext, audit metadata или browser storage."),
    ("Testing style", "работа строилась через failing tests, browser evidence, DB assertions, secret scan и regression guards."),
])

doc.add_heading("3. Хронология работы и история отчётов в чате", level=1)
doc.add_paragraph("Ниже приведена реконструкция того, что было зафиксировано в моих промежуточных и финальных отчётах внутри этого чата. Это не дословный transcript, а инженерная сводка по доступной истории сообщений, планам, финальным статусам и выполненным командам.")
table(doc, ["Этап", "Что было сказано/зафиксировано", "Инженерный смысл"], report_history, [1.7, 3.2, 3.0])

doc.add_heading("4. Изменения по доменам", level=1)
domain_sections = [
    ("4.1 Backend AI Gateway / CometAPI / Chat", [
        "Добавлен отдельный live chat streaming path, отделённый от structured automation JSON route.",
        "Streaming request собирается backend-only: POST /chat/completions, stream=true, max_tokens=256, reasoning_effort=high, thinking enabled.",
        "SSE parser расширен для reasoning-only/content/mixed chunks; visible assistant content аккумулируется отдельно от hidden reasoning.",
        "Provider-specific fields не отбрасываются DTO/SDK уровнем; для CometAPI применяется backend-only fetch path.",
        "Happy path не должен использовать canned/Stage18 text; assistant message сохраняется только после реального provider content.",
        "Provider failure пишет controlled stream_failed event/audit без raw provider error/key/header.",
    ]),
    ("4.2 Settings / AI route readiness", [
        "Connection test больше не считает /models достаточным PASS для chat route.",
        "Для CometAPI/chat route readiness должен подтверждаться chat completion visible-content probe.",
        "AI key remains write-only: сохранённый ключ не возвращается UI; используются fingerprint/status/metadata.",
        "Route preference и secret ref обновляются через backend intake path, а не plaintext DB patch.",
    ]),
    ("4.3 Project chat UI/sidebar", [
        "ProjectSidebar переведён на project-scoped chat list вместо one-shot searchChats для блока Чаты.",
        "После stream success/failure инвалидируются project chat queries/snapshot, чтобы history/sidebar обновлялись без ручного reload.",
        "Assistant marker/search scoped to assistant bubble, чтобы user prompt не считался provider response.",
        "Global/floating composer не дублируется на project chat route.",
    ]),
    ("4.4 Automation runtime/canvas", [
        "Activepieces iframe route открывается и сохраняет project/workspace context.",
        "Добавлен LexFrame overlay button Запустить dry-run на /automation route.",
        "Run стартует через LexFrame backend startAutomationRun, не через прямой browser call к внешнему runtime.",
        "DB пишет workflow_run, activepieces_run_binding, runtime token hash/JTI evidence и audit events.",
        "Canvas branch test reached backend; controlled failure для текущего workflow without branch runtime mapping.",
    ]),
    ("4.5 File upload", [
        "UploadDialog получил настоящий input type=file.",
        "File.name/type/size берутся из browser File object.",
        "Content отправляется на backend as base64 bytes до complete upload.",
        "Backend считает sha256, сверяет size/hash/mime, пишет audit safe metadata.",
        "Complete upload получает sha256 из backend content response и создаёт document/version/storage evidence.",
    ]),
    ("4.6 Secret-safety", [
        "Найден и исправлен Activepieces JWT leak в browser sessionStorage после iframe route.",
        "Добавлен shared cleanup helper для Activepieces browser session tokens.",
        "AppShell очищает stale Activepieces tokens вне /automation embed route.",
        "Artifacts with old JWT/signed-url-like hits были redacted; повторный scan дал нули.",
    ]),
]
for title, items in domain_sections:
    doc.add_heading(title, level=2)
    bullets(doc, items)

doc.add_heading("5. Подробный файл-за-файлом обзор", level=1)
doc.add_paragraph("Таблица отражает текущий git diff на момент подготовки документа. Некоторые изменения в AI/chat/settings были сделаны в предыдущих проходах этого же чата, а изменения automation/upload/token cleanup — в последнем runtime QA проходе.")
table(doc, ["Область", "Файл", "Изменение", "Назначение"], tracked_files, [1.5, 2.9, 0.9, 3.0])

doc.add_heading("6. Баги, root cause и fixes", level=1)
table(doc, ["Баг", "Симптом", "Root cause", "Fix", "Evidence / результат"], bugs, [1.5, 1.8, 2.0, 2.0, 1.7])

doc.add_heading("7. Runtime evidence: browser, DB, audit", level=1)
doc.add_paragraph("Ниже перечислены факты, полученные из Playwright UI-smoke и DB queries. Значения raw tokens/keys/headers намеренно не выводились и не попали в документ.")
table(doc, ["Домен", "Источник", "Идентификатор / строка", "Evidence"], evidence, [1.6, 2.0, 2.3, 2.6])
doc.add_heading("7.1 Browser evidence", level=2)
bullets(doc, [
    ("Automation iframe", "Route /app/projects/project_claim_001/automations/16000000-0000-4000-8000-000000008001/automation loaded Activepieces builder iframe."),
    ("Dry-run button", "Button Запустить dry-run visible and clickable; UI reported Запуск создан with run id."),
    ("Upload dialog", "Documents page: New upload -> Select file -> Upload selected file; flow completed for document id da52868e..."),
    ("Chat branch", "Sidebar shows project chats including branch entry Новый чат проекта / ветка; route opened successfully."),
    ("Canvas branch test", "Branch mode + selected manual_start + Run produced controlled diagnostic in UI: Runtime mapping is missing."),
    ("Direct external requests", "Browser network filter for api.cometapi.com / activepieces external hosts showed 0 direct provider/runtime external requests in the checked session."),
    ("Storage after fix", "After leaving Activepieces iframe route, localStorage/sessionStorage/cookies secret-pattern hits = 0."),
])
doc.add_heading("7.2 DB evidence queries summary", level=2)
code(doc, """workflow_runs checked: 2
latest automation dry-run: 45ec0f45-d43d-4c9f-aada-3ab095161692
uploaded document checked: da52868e-d0b1-46ba-8f9f-4426b95924f5
chat branch rows checked: 1
canvas test run rows checked: 1
raw_secret_pattern_hits over checked DB text/json evidence: 0""")

doc.add_heading("8. Команды и результаты", level=1)
table(doc, ["Команда", "Результат", "Что доказано"], commands, [3.6, 1.2, 3.2])

doc.add_heading("9. Secret-safety и redaction", level=1)
bullets(doc, [
    ("Repo secret scan", "corepack pnpm secret-scan -> PASS."),
    ("Runtime artifacts", "Скан artifacts и .playwright-mcp после redaction: providerKey=0, authBearer=0, bearerSk=0, jwt=0, signedUrl=0, privateKey=0."),
    ("DB scan", "Checked audit metadata, documents, document_versions, workflow_runs, activepieces bindings/sessions: raw_secret_pattern_hits=0."),
    ("Browser storage", "После ухода с iframe route Activepieces token absent; local/session/cookie secret-pattern hits = 0."),
    ("Network", "Нет прямых browser calls к api.cometapi.com или внешним Activepieces/API hosts в проверенном session; runtime route идёт через LexFrame /automation-runtime или /api endpoints."),
    ("Artifact cleanup", "Старый stage17 JWT-like artifact и stage19 signed-url-like artifact были redacted, после чего повторный artifact scan дал нули."),
])
note(doc, "Секреты в этом отчёте", "Документ не содержит raw ключей, JWT, Authorization headers или signed URLs. Даже при описании credential work используются только fingerprints, lengths и safe status.", "E2F0D9")

doc.add_heading("10. Оставшиеся ограничения и следующий шаг", level=1)
bullets(doc, [
    ("Canvas branch live execution", "Текущий workflow не содержит condition/router branch node. Branch test пишет DB/audit, но status failed/invalid с RUNTIME_MAPPING_MISSING. Для полного PASS нужен disposable branch-capable workflow fixture или UI-flow добавления condition router."),
    ("File content persistence semantics", "В local dev pass backend принимает bytes и считает sha256, а complete flow создаёт storage metadata. Если product requirement требует физическую запись object bytes в external storage, следующий шаг — подключить signed/direct storage write path вместо local metadata completion."),
    ("E2E file", "В этом проходе browser smoke выполнен вручную через Playwright tool plus tests/typechecks. Можно дополнить отдельным tests/e2e/stage21-runtime-system-smoke.spec.ts, если нужен формальный Playwright test artifact на CI."),
    ("Dirty worktree", "Рабочая копия содержит незакоммиченные изменения из нескольких проходов этого чата. Перед merge/push нужен осознанный commit scope, чтобы не смешать unrelated artifacts."),
])

doc.add_heading("11. Appendix A — status matrix", level=1)
table(doc, ["Домен", "Статус", "Комментарий"], [
    ("Settings", "PASS", "Targeted settings tests/typechecks previously passed; not rewritten in last runtime QA except AI readiness improvements from earlier pass."),
    ("Projects", "PASS", "Project context preserved in tested routes; project_claim_001 used for chat/automation/upload evidence."),
    ("Chat + AI provider", "Targeted tests PASS; live provider smoke not rerun in final runtime-upload pass", "Backend chat service and LexFrameChatShell tests pass; earlier chat/provider work changed streaming path."),
    ("Automations + runtime/canvas", "PASS for iframe/open/run; PARTIAL for live canvas branch execution", "Dry-run created DB/audit evidence; branch test controlled failed due missing branch node mapping."),
    ("File upload", "PASS", "Real file selected and bytes sent to backend, sha256 stored, DB/audit evidence exists."),
    ("Secret-safety", "PASS after redaction", "Repo, DB checked evidence, browser storage and runtime artifact pattern scans clean after cleanup/redaction."),
], [1.8, 2.0, 4.2])

doc.add_heading("Appendix B — detailed checklists", level=1)
for title, items in {
    "Automation runtime checklist": [
        "Entry point автоматизаций видим и кликабелен.",
        "Active project/workspace context preserved: project_claim_001.",
        "Activepieces canvas iframe opens and shows builder, not raw placeholder.",
        "LexFrame dry-run button visible on embed route.",
        "Dry-run creates workflow_run and activepieces_run_binding.",
        "Runtime token evidence stored as hash/JTI only.",
        "Audit events written without raw tokens.",
        "Leaving iframe clears Activepieces browser token.",
    ],
    "File upload checklist": [
        "Upload dialog exposes actual file input.",
        "File name/type/size derive from selected File object.",
        "Bytes sent to backend content endpoint before complete.",
        "Backend validates canonical base64, size, MIME and optional sha256.",
        "Backend computes sha256 and returns it to frontend.",
        "Complete upload uses backend sha256.",
        "Document/version/storage rows exist.",
        "Audit contains safe metadata, not raw bytes.",
    ],
    "Chat/sidebar checklist": [
        "Sidebar project chat list uses project scoped data.",
        "Sidebar shows existing chats and branch entries.",
        "Stream success/failure invalidates project chat snapshots.",
        "Assistant marker is scoped to assistant message in tests.",
        "Frontend does not call provider directly.",
    ],
    "Secret-safety checklist": [
        "No provider key in repo scan.",
        "No Authorization Bearer in artifacts scan.",
        "No JWT shape in artifacts scan after redaction.",
        "No signed URL pattern in artifacts scan after redaction.",
        "No raw secret hits in checked DB/audit text/json fields.",
        "Browser storage scan after iframe route cleanup is zero.",
    ],
}.items():
    doc.add_heading(title, level=2)
    bullets(doc, items)

doc.add_heading("Appendix C — notes on why the fixes are minimal", level=1)
for paragraph in [
    "Automation run fix is local to the Activepieces iframe route. It does not rewrite Activepieces integration, session issuance, readiness checks, or runtime token model. It simply exposes an existing backend run capability in the route where the user actually works.",
    "Upload fix is additive and follows the existing upload intent/complete contract. It adds a content intake step and safe validation without introducing frontend storage secrets or direct external object storage calls from the browser.",
    "Activepieces token cleanup is targeted to tokens with issuer activepieces and only keys known to be used by the embedded runtime. It does not wipe unrelated app auth/session storage.",
    "Chat/provider changes keep provider calls backend-only and do not switch automation AI to chat model. Structured automation route remains separate from normal chat completions.",
    "Artifact redaction changed only legacy evidence files containing secret-pattern-like artifacts. It did not alter source code behavior except to ensure deliverable safety and secret-scan compliance.",
]:
    doc.add_paragraph(paragraph)

doc.add_heading("Appendix D — final operational status", level=1)
table(doc, ["Question", "Answer"], [
    ("Можно ли открыть automation canvas?", "Да, iframe route opens Activepieces builder."),
    ("Можно ли запустить automation?", "Да, кнопка Запустить dry-run создаёт workflow_run через LexFrame backend."),
    ("Работают ли ветки чата?", "Существующая chat branch persisted and visible in sidebar; route opens."),
    ("Работает ли canvas branch live test?", "Controlled partial: backend test-run records/audit exist, but current workflow lacks branch runtime mapping."),
    ("Можно ли загрузить файл?", "Да, real file selected, bytes sent to backend, sha256 computed, document/version/storage rows created."),
    ("Есть ли direct browser provider call?", "В проверенном browser session — 0."),
    ("Есть ли raw secret/JWT/Auth hits?", "После cleanup/redaction — 0 in checked repo/artifact/browser/DB evidence scans."),
], [2.8, 5.0])

# Repeated appendix-style detail to satisfy the user's request for a large, maximally detailed document without adding unsafe raw secrets.
doc.add_heading("Appendix E — expanded narrative log", level=1)
for i, (title, items) in enumerate(domain_sections, start=1):
    doc.add_heading(f"E.{i} {title}", level=2)
    for item in items:
        doc.add_paragraph(item)
    doc.add_paragraph("Engineering note: this change was kept inside the smallest layer that owned the broken behavior. Tests and runtime evidence were used to avoid a visual-only acceptance.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("End of report — generated by Codex for LexFrame runtime QA/fix chat")
r.italic = True
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(90, 90, 90)

props = doc.core_properties
props.title = "LexFrame latest changes report"
props.subject = "Changes performed during Codex chat"
props.author = "Codex"
props.keywords = "LexFrame, Stage21, automation, chat, upload, audit, secret-safety"
props.comments = "No raw secrets included."
props.created = datetime(2026, 5, 10, 12, 0, 0)

doc.save(OUT)
print(OUT)
